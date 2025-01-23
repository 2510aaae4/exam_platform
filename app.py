from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_from_directory, abort, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
import os
import json
from datetime import datetime
from functools import wraps
import traceback
from config import Config
from docx import Document
from docx.shared import Pt
import io

app = Flask(__name__)
app.config.from_object(Config)

# 初始化 SQL Database
db = SQLAlchemy(app)
migrate = Migrate(app, db)



# 資料庫模型
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    exam_records = db.relationship('ExamRecord', backref='user', lazy=True)
    notes = db.relationship('Note', backref='user', lazy=True)

class ExamRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    exam_id = db.Column(db.String(50), nullable=False)
    date = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    score = db.Column(db.Float, nullable=False)
    answers = db.Column(db.JSON, nullable=False)  # 儲存用戶答案
    confidence = db.Column(db.JSON, nullable=True)  # 儲存把握度
    exam_duration = db.Column(db.Integer, nullable=True)  # 儲存考試時間（秒）
    notes = db.relationship('Note', backref='exam_record', lazy=True)

    def to_dict(self):
        return {
            'id': self.id,
            'exam_id': self.exam_id,
            'date': self.date.strftime('%Y-%m-%d %H:%M:%S'),
            'score': self.score,
            'answers': self.answers,
            'confidence': self.confidence,
            'exam_duration': self.exam_duration
        }

class Note(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    exam_record_id = db.Column(db.Integer, db.ForeignKey('exam_record.id'), nullable=False)
    question_number = db.Column(db.Integer, nullable=False)
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    __table_args__ = (
        db.Index('idx_note_exam_question', 'exam_record_id', 'question_number'),
    )

    def to_dict(self):
        return {
            'id': self.id,
            'user_id': self.user_id,
            'exam_record_id': self.exam_record_id,
            'question_number': self.question_number,
            'content': self.content,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': self.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        }

class Favorite(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    exam_id = db.Column(db.String(50), nullable=False)
    question_number = db.Column(db.Integer, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    __table_args__ = (
        db.UniqueConstraint('user_id', 'exam_id', 'question_number', name='unique_user_question_favorite'),
    )
    
    def to_dict(self):
        return {
            'id': self.id,
            'exam_id': self.exam_id,
            'question_number': self.question_number,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S')
        }

# 只在首次運行時創建表
with app.app_context():
    try:
        # 嘗試查詢用戶表
        try:
            User.query.first()
        except:
            # 如果查詢失敗，表示表不存在，創建所有表
            db.create_all()
            print("Database tables created successfully")
            
            # 從 allowed_users.txt 讀取並創建用戶
            with open('allowed_users.txt', 'r') as f:
                allowed_users = f.read().splitlines()
                for username in allowed_users:
                    if username.strip():
                        user = User(username=username.strip())
                        db.session.add(user)
                db.session.commit()
            print("Initial users created successfully")
    except Exception as e:
        print(f"Error initializing database: {str(e)}")
        print(traceback.format_exc())

# 驗證用戶是否存在
def validate_user(username):
    with open('allowed_users.txt', 'r') as f:
        allowed_users = f.read().splitlines()
    return username in allowed_users

def handle_errors(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except Exception as e:
            print(f"Error in {f.__name__}:")
            print(traceback.format_exc())
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500
    return wrapper

def load_exam_content(exam_id):
    try:
        year = exam_id.split('_')[0]
        part = exam_id.split('_')[1]
        
        base_dir = os.path.dirname(__file__)
        print(f"Base directory: {base_dir}")
        
        # 讀取題目文件
        questions_file = os.path.join(base_dir, 'questions', year, f'{year}_{part}.md')
        answers_file = os.path.join(base_dir, 'questions', year, f'{year}_{part}_answers.txt')
        
        print(f"Attempting to load exam content from: {questions_file}")
        print(f"Answers file path: {answers_file}")
        
        if not os.path.exists(questions_file):
            print(f"Questions file not found: {questions_file}")
            return None
            
        if not os.path.exists(answers_file):
            print(f"Answers file not found: {answers_file}")
            return None
            
        # 讀取答案文件
        with open(answers_file, 'r', encoding='utf-8-sig') as f:
            answers_content = f.read().strip().split('\n')
            answers = []
            for line in answers_content:
                if line.strip() in ['A', 'B', 'C', 'D', 'E']:
                    # 將答案轉換為數字索引 (A=0, B=1, C=2, D=3, E=4)
                    answers.append(ord(line.strip()) - ord('A'))
        
        # 使用 utf-8-sig 來處理可能的 BOM
        with open(questions_file, 'r', encoding='utf-8-sig') as f:
            content = f.read()
            
        print(f"Successfully read questions file. Content length: {len(content)}")
            
        questions = []
        current_question = None
        current_content = []
        
        # 預處理內容，移除特殊字符
        content = content.replace('\ufeff', '')  # 移除 BOM
        content = ''.join(char for char in content if ord(char) < 65536)  # 只保留基本多語言平面字符
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:  # 跳過空行
                continue
                
            if line.startswith('## '):  # 新題目開始
                if current_question:
                    if current_content:
                        current_question['content'] = '\n'.join(current_content)
                    questions.append(current_question)
                question_text = line[3:].strip()  # 移除 "## "
                # 從題目文字中提取題號
                question_number = question_text.split('.')[0]
                current_question = {
                    'number': question_number,
                    'content': question_text,
                    'options': [],
                    'image': None
                }
                # 檢查是否有對應的圖片
                image_path = os.path.join(base_dir, 'questions', year, 'image', f"{question_number}.png")
                print(f"Checking for image: {image_path}")
                if os.path.exists(image_path):
                    relative_path = os.path.join(year, 'image', f"{question_number}.png")
                    image_url = url_for('serve_questions', filename=relative_path, _external=True)
                    current_question['image'] = image_url
                    print(f"Found image for question {question_number}. URL: {image_url}")
                else:
                    print(f"No image found for question {question_number}")
                current_content = [question_text]
            elif line.startswith(('A.', 'B.', 'C.', 'D.', 'E.')):  # 選項
                if current_question:
                    current_question['options'].append(line[2:].strip())
            else:
                if current_question and not line.startswith('##'):
                    current_content.append(line)
                    
        # 處理最後一題
        if current_question:
            if current_content:
                current_question['content'] = '\n'.join(current_content)
            questions.append(current_question)
            
        return {
            'questions': questions,
            'answers': answers
        }
        
    except Exception as e:
        print(f"Error loading exam content: {str(e)}")
        print(traceback.format_exc())
        return None

def convert_letter_to_index(letter):
    """將答案字母轉換為索引"""
    mapping = {'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4}
    return mapping.get(letter.strip().upper())

def calculate_score(exam_id, user_answers):
    """計算考試分數並返回正確答案"""
    try:
        # 解析年份和部分
        year = exam_id.split('_')[0]
        part = exam_id.split('_')[1]
        
        # 構建答案文件路徑
        answer_file = os.path.join('questions', year, f'{year}_{part}_answers.txt')
        
        if not os.path.exists(answer_file):
            raise Exception('找不到答案文件')
        
        # 讀取並轉換答案
        with open(answer_file, 'r', encoding='utf-8') as f:
            correct_answers = []
            for line in f:
                letter = line.strip()
                if letter:
                    index = convert_letter_to_index(letter)
                    if index is not None:
                        correct_answers.append(index)
                    else:
                        raise Exception('答案文件格式錯誤')
        
        if len(correct_answers) != 80:
            raise Exception('答案文件格式錯誤')
        
        # 計算分數
        score = 0
        correct_question_numbers = []  # 記錄答對的題號
        for i, answer in enumerate(user_answers):
            if answer is not None and answer == correct_answers[i]:
                score += 1.25
                correct_question_numbers.append(str(i + 1))  # 題號從1開始
        
        return {
            'score': score,
            'correct_answers': correct_question_numbers
        }
        
    except Exception as e:
        print(f"Error calculating score: {str(e)}")
        raise e

# 從考試ID獲取考試標題
def get_exam_title(exam_id):
    try:
        year = exam_id.split('_')[0]
        part = '前80題' if exam_id.endswith('_1') else '後80題'
        return f'{year}年度{part}'
    except Exception as e:
        print(f"Error generating exam title for {exam_id}: {str(e)}")
        return exam_id

@app.route('/')
def index():
    if not session.get('username'):
        return render_template('login.html')
    if session.get('is_admin'):
        return render_template('admin_index.html')
    return render_template('index.html')

@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username')
    if not username:
        return jsonify({'success': False, 'message': '請輸入用戶名'})
    
    if username == 'admin_107492':
        session['username'] = username
        session['is_admin'] = True
        return jsonify({'success': True, 'is_admin': True})
    
    if validate_user(username):
        session['username'] = username
        session['is_admin'] = False
        return jsonify({'success': True, 'is_admin': False})
    
    return jsonify({'success': False, 'message': '無效的用戶名'})

@app.route('/simulation')
def simulation():
    if 'username' not in session:
        return redirect(url_for('index'))
    return render_template('simulation.html')

@app.route('/past_exams')
def past_exams():
    if 'username' not in session:
        return redirect(url_for('index'))
    return render_template('past_exams.html')

@app.route('/load-exam/<exam_id>')
@handle_errors
def load_exam_route(exam_id):
    if 'username' not in session:
        return jsonify({
            'success': False,
            'error': '請先登入'
        }), 401
    
    print(f"User {session['username']} attempting to load exam {exam_id}")
    
    try:
        # 使用 load_exam_content 函數載入考卷
        exam_content = load_exam_content(exam_id)
        
        if exam_content is None:
            return jsonify({
                'success': False,
                'error': '找不到試卷文件'
            }), 404
            
        return jsonify({
            'success': True,
            'questions': exam_content['questions']
        })
        
    except Exception as e:
        print(f"Error loading exam: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'載入試卷時發生錯誤：{str(e)}'
        }), 500

@app.route('/submit-exam', methods=['POST'])
@handle_errors
def submit_exam():
    try:
        data = request.get_json()
        print('Received exam data:', data)  # 添加日誌
        
        exam_id = data.get('examId')
        answers = data.get('answers')
        confidence = data.get('confidence')
        exam_duration = data.get('examDuration', 0)  # 新增：接收考試時間
        
        print('Exam duration:', exam_duration)  # 添加日誌
        
        # 驗證必要參數
        if not exam_id or not answers:
            return jsonify({
                'success': False,
                'error': '缺少必要資料'
            })
            
        # 計算分數
        score_result = calculate_score(exam_id, answers)
        score = score_result['score']
        correct_answers = score_result['correct_answers']
        
        # 獲取當前用戶
        username = session.get('username')
        if not username:
            return jsonify({
                'success': False,
                'error': '請先登入'
            })
            
        user = User.query.filter_by(username=username).first()
        if not user:
            return jsonify({
                'success': False,
                'error': '用戶不存在'
            })
            
        # 創建考試記錄
        exam_record = ExamRecord(
            user_id=user.id,
            exam_id=exam_id,
            score=score,
            answers=answers,
            confidence=confidence,
            exam_duration=exam_duration  # 新增：儲存考試時間
        )
        
        db.session.add(exam_record)
        db.session.commit()
        
        print('Saved exam record with duration:', exam_record.exam_duration)  # 添加日誌
        
        return jsonify({
            'success': True,
            'score': score,
            'correct_answers': correct_answers,
            'recordId': exam_record.id
        })
    except Exception as e:
        print('Error in submit_exam:', str(e))  # 添加日誌
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/score-result/<int:record_id>')
def score_result(record_id):
    if 'username' not in session:
        return redirect(url_for('index'))

    try:
        record = ExamRecord.query.get(record_id)
        if not record:
            abort(404)

        user = User.query.filter_by(username=session['username']).first()
        if not user or record.user_id != user.id:
            abort(403)

        return render_template('score_result.html', record=record)

    except Exception as e:
        print(f"Error in score_result: {str(e)}")
        return render_template('404.html'), 404

@app.route('/api/exam-notes/<record_id>')
def get_exam_notes(record_id):
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'})

    try:
        user = User.query.filter_by(username=session['username']).first()
        if not user:
            return jsonify({'success': False, 'error': '找不到用戶'})

        # 獲取特定考試記錄
        exam_record = ExamRecord.query.get(record_id)
        if not exam_record:
            return jsonify({'success': False, 'error': '找不到考試記錄'})

        # 只獲取當前用戶對這個考試記錄的筆記
        notes = Note.query.filter_by(
            user_id=user.id,
            exam_record_id=record_id
        ).order_by(Note.question_number).all()

        # 按題號組織筆記
        notes_by_question = {}
        for note in notes:
            notes_by_question[str(note.question_number)] = [{
                'content': note.content,
                'created_at': note.created_at.strftime('%Y-%m-%d %H:%M:%S')
            }]

        return jsonify({
            'success': True,
            'notes': notes_by_question
        })

    except Exception as e:
        print("Error:", str(e))
        return jsonify({'success': False, 'error': str(e)})

@app.route('/save-notes', methods=['POST'])
def save_notes():
    if 'username' not in session:
        return jsonify({'success': False, 'message': '請先登入'})

    try:
        data = request.get_json()
        record_id = data.get('record_id')
        notes_data = data.get('notes')
        
        print(f"record_id: {record_id}, notes_data: {notes_data}")  # 調試輸出
        
        if not record_id or notes_data is None:
            return jsonify({'success': False, 'message': '缺少必要參數'})
            
        user = User.query.filter_by(username=session['username']).first()
        if not user:
            return jsonify({'success': False, 'message': '用戶不存在'})

        exam_record = ExamRecord.query.get(record_id)
        if not exam_record:
            return jsonify({'success': False, 'message': '考試記錄不存在'})

        # 刪除該用戶對該考試的所有舊筆記
        Note.query.filter_by(
            user_id=user.id,
            exam_record_id=record_id
        ).delete()

        # 添加新筆記
        for question_num, content in notes_data.items():
            if content.strip():  # 只保存非空筆記
                note = Note(
                    user_id=user.id,
                    exam_record_id=record_id,
                    question_number=int(question_num),
                    content=content.strip()
                )
                db.session.add(note)

        db.session.commit()
        return jsonify({'success': True})

    except Exception as e:
        db.session.rollback()
        print(f"Error in save_notes: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/exam-records')
def get_exam_records():
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'})

    try:
        user = User.query.filter_by(username=session['username']).first()
        if not user:
            return jsonify({'success': False, 'error': '找不到使用者'})

        # 獲取所有考試記錄並按考試ID分組
        records = ExamRecord.query.filter_by(user_id=user.id).order_by(ExamRecord.date.desc()).all()
        grouped_records = {}

        for record in records:
            exam_id = record.exam_id
            if exam_id not in grouped_records:
                grouped_records[exam_id] = {
                    'title': get_exam_title(exam_id),
                    'highest_score': record.score,
                    'attempts': []
                }
            else:
                if record.score > grouped_records[exam_id]['highest_score']:
                    grouped_records[exam_id]['highest_score'] = record.score

            # 計算各科答對率
            subject_mapping = get_subject_mapping(exam_id)
            answers = record.answers
            score_result = calculate_score(exam_id, answers)
            correct_answers = score_result['correct_answers']
            
            # 初始化各科統計，只包含這次考試有考的科別
            subject_stats = {}
            # 根據考試ID判斷題號範圍
            start_question = 1 if exam_id.endswith('_1') else 81
            end_question = 80 if exam_id.endswith('_1') else 160
            
            # 只處理該考卷範圍內的題目
            for q_num in range(start_question, end_question + 1):
                q_str = str(q_num)
                if q_str in subject_mapping:  # 確保題號有對應的科別
                    subject = subject_mapping[q_str]
                    if subject not in subject_stats:
                        subject_stats[subject] = {'correct': 0, 'total': 0}
                    subject_stats[subject]['total'] += 1
                    if q_str in correct_answers:
                        subject_stats[subject]['correct'] += 1
            
            # 計算各科答對率，只包含有考的科別
            subject_percentages = {}
            for subject, stats in subject_stats.items():
                if stats['total'] > 0:  # 只包含有考的科別
                    subject_percentages[subject] = (stats['correct'] / stats['total']) * 100

            # 添加考試記錄
            attempt = {
                'record_id': record.id,
                'date': record.date.strftime('%Y-%m-%d %H:%M:%S'),
                'score': record.score,
                'exam_duration': record.exam_duration,
                'subject_percentages': subject_percentages
            }
            grouped_records[exam_id]['attempts'].append(attempt)

        # 轉換為列表格式
        records_list = []
        for exam_id, data in grouped_records.items():
            records_list.append({
                'exam_id': exam_id,
                'title': data['title'],
                'highest_score': data['highest_score'],
                'attempts': sorted(data['attempts'], key=lambda x: x['date'], reverse=True)
            })

        return jsonify({
            'success': True,
            'records': sorted(records_list, key=lambda x: x['exam_id'])
        })
    except Exception as e:
        print(f"Error in get_exam_records: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)})

@app.route('/exam-detail/<int:record_id>')
def exam_detail(record_id):
    """顯示考試詳細資訊的頁面"""
    try:
        if 'username' not in session:
            return redirect(url_for('login'))
            
        # 檢查考試記錄是否存在
        record = ExamRecord.query.get(record_id)
        if not record:
            flash('找不到考試記錄')
            return redirect(url_for('dashboard'))
            
        # 檢查權限
        current_user = User.query.filter_by(username=session['username']).first()
        if record.user_id != current_user.id:
            flash('沒有權限查看此考試記錄')
            return redirect(url_for('dashboard'))
            
        return render_template('exam_detail.html',
                             username=session['username'],
                             record_id=record_id)
                             
    except Exception as e:
        print(f"Error in exam_detail: {str(e)}")
        print(traceback.format_exc())
        return render_template('exam_detail.html',
                             username=session['username'],
                             record_id=record_id,
                             error="載入考試詳情時發生錯誤")

@app.route('/api/save-notes', methods=['POST'])
def save_notes_api():
    """儲存筆記的API端點"""
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'})
        
    try:
        data = request.get_json()
        print(f"Received notes data: {data}")  # 調試輸出
        
        if not data:
            return jsonify({'success': False, 'error': '沒有收到數據'})
            
        record_id = data.get('record_id')
        notes_data = data.get('notes')
        
        print(f"record_id: {record_id}, notes_data: {notes_data}")  # 調試輸出
        
        if not record_id or notes_data is None:
            return jsonify({'success': False, 'error': '缺少必要參數'})
            
        user = User.query.filter_by(username=session['username']).first()
        if not user:
            return jsonify({'success': False, 'error': '用戶不存在'})

        exam_record = ExamRecord.query.get(record_id)
        if not exam_record:
            return jsonify({'success': False, 'error': '考試記錄不存在'})

        # 檢查權限
        user_id = User.query.filter_by(username=session['username']).first().id
        if exam_record.user_id != user_id:
            return jsonify({'success': False, 'error': '沒有權限'})
            
        # 更新或創建筆記
        for question_number, content in notes_data.items():
            # 如果 content 是字典，獲取其內容
            if isinstance(content, dict):
                content = content.get('content', '')
            
            # 查找現有筆記
            note = Note.query.filter_by(
                exam_record_id=record_id,
                question_number=question_number
            ).first()
            
            if note:
                # 更新現有筆記
                note.content = content
                note.updated_at = datetime.utcnow()
            else:
                # 創建新筆記
                note = Note(
                    user_id=user_id,
                    exam_record_id=record_id,
                    question_number=question_number,
                    content=content
                )
                db.session.add(note)
                
        db.session.commit()
        return jsonify({'success': True})
        
    except Exception as e:
        db.session.rollback()
        print(f"Error saving notes: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/get-exam-notes/<int:record_id>')
def get_exam_notes_api(record_id):
    """獲取考試筆記的API端點"""
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'})
        
    try:
        # 獲取考試記錄
        record = ExamRecord.query.get_or_404(record_id)
        
        # 檢查權限
        if record.user_id != User.query.filter_by(username=session['username']).first().id:
            return jsonify({'success': False, 'error': '沒有權限'})
            
        # 獲取所有筆記
        notes = Note.query.filter_by(exam_record_id=record_id).all()
        
        # 將筆記組織成以題號為鍵的字典
        notes_dict = {}
        for note in notes:
            if note.question_number not in notes_dict:
                notes_dict[note.question_number] = []
            notes_dict[note.question_number].append({
                'content': note.content,
                'created_at': note.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': note.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            })
            
        return jsonify({
            'success': True,
            'notes': notes_dict
        })
        
    except Exception as e:
        print(f"Error getting notes: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)})

def get_subject_mapping(exam_id):
    """獲取題目對應的科別映射"""
    try:
        year = exam_id.split('_')[0]
        
        # 讀取 d1 和 d2 檔案
        base_dir = os.path.dirname(__file__)
        d1_file = os.path.join(base_dir, 'questions', year, f'{year}_d1.md')
        d2_file = os.path.join(base_dir, 'questions', year, f'{year}_d2.md')
        
        print(f"Looking for d1 file: {d1_file}")
        print(f"Looking for d2 file: {d2_file}")
        
        subject_mapping = {}
        
        # 處理 d1 檔案 (1-80題)
        if os.path.exists(d1_file):
            print(f"Found d1 file")
            with open(d1_file, 'r', encoding='utf-8-sig') as f:
                lines = f.readlines()
                print(f"D1 content: {lines[:10]}")  # 顯示前10行
                
                # 每行對應一題
                for i, line in enumerate(lines, start=1):
                    subject = line.strip()
                    if subject:  # 如果不是空行
                        subject_mapping[str(i)] = subject
                        print(f"Mapped question {i} to subject: {subject}")
                    if i >= 80:  # 只處理到第80題
                        break
        
        # 處理 d2 檔案 (81-160題)
        if os.path.exists(d2_file):
            print(f"Found d2 file")
            with open(d2_file, 'r', encoding='utf-8-sig') as f:
                lines = f.readlines()
                print(f"D2 content: {lines[:10]}")  # 顯示前10行
                
                # 每行對應一題，從第81題開始
                for i, line in enumerate(lines, start=81):
                    subject = line.strip()
                    if subject:  # 如果不是空行
                        subject_mapping[str(i)] = subject
                        print(f"Mapped question {i} to subject: {subject}")
                    if i >= 160:  # 只處理到第160題
                        break
        
        print(f"Final subject mapping: {subject_mapping}")
        return subject_mapping
        
    except Exception as e:
        print(f"Error in get_subject_mapping: {str(e)}")
        print(traceback.format_exc())
        return {}

@app.route('/api/exam-detail/<int:record_id>')
def get_exam_detail(record_id):
    """獲取考試詳細資訊的API端點"""
    try:
        if 'username' not in session:
            return jsonify({'success': False, 'error': '請先登入'})
            
        record = ExamRecord.query.get(record_id)
        if not record:
            return jsonify({'success': False, 'error': '找不到考試記錄'})
            
        # 檢查權限
        current_user = User.query.filter_by(username=session['username']).first()
        if record.user_id != current_user.id:
            return jsonify({'success': False, 'error': '沒有權限查看此考試記錄'})
            
        # 載入考試內容
        exam_content = load_exam_content(record.exam_id)
        if not exam_content:
            return jsonify({'success': False, 'error': '無法載入考試內容'})
            
        # 獲取考試科目映射
        subjects = get_subject_mapping(record.exam_id)
        print(f"Exam ID: {record.exam_id}")
        print(f"Loaded subjects: {subjects}")  # 調試輸出
        
        # 解析用戶答案
        user_answers = {}
        try:
            if isinstance(record.answers, str):
                user_answers = json.loads(record.answers)
            elif isinstance(record.answers, dict):
                user_answers = record.answers
            elif isinstance(record.answers, list):
                user_answers = {str(i): ans for i, ans in enumerate(record.answers)}
        except Exception as e:
            print(f"Error parsing answers: {e}")
            print(f"Original answers: {record.answers}")
            
        print(f"Parsed user answers: {user_answers}")  # 調試輸出
        
        # 解析信心程度
        confidence_levels = {}
        if record.confidence:
            try:
                if isinstance(record.confidence, str):
                    confidence_levels = json.loads(record.confidence)
                elif isinstance(record.confidence, dict):
                    confidence_levels = record.confidence
                elif isinstance(record.confidence, list):
                    confidence_levels = {str(i): conf for i, conf in enumerate(record.confidence)}
            except Exception as e:
                print(f"Error parsing confidence: {e}")
                confidence_levels = {}
                
        # 準備問題列表
        questions = []
        for i, question in enumerate(exam_content['questions']):
            question_number = question['number']
            correct_answer = exam_content['answers'][i] if i < len(exam_content['answers']) else None
            
            # 獲取用戶答案（嘗試多種可能的鍵）
            user_answer = None
            possible_keys = [str(i), str(question_number), i, question_number]
            for key in possible_keys:
                if str(key) in user_answers:
                    user_answer = user_answers[str(key)]
                    break
                    
            print(f"Question {question_number}: user_answer = {user_answer}, correct_answer = {correct_answer}")  # 調試輸出
            
            # 獲取信心程度
            confidence = 'medium'  # 默認值
            for key in possible_keys:
                if str(key) in confidence_levels:
                    confidence = confidence_levels[str(key)]
                    break
                    
            # 獲取科目
            subject = subjects.get(str(question_number), '未分類')
            print(f"Question {question_number} subject: {subject}")  # 調試輸出
            
            question_data = {
                'number': question_number,
                'content': question['content'],
                'options': question['options'],
                'user_answer': user_answer,  # 用戶的答案
                'correct_answer': correct_answer,  # 正確答案
                'is_correct': user_answer == correct_answer if user_answer is not None else None,
                'confidence': confidence,
                'subject': subject,
                'image': question.get('image')
            }
            questions.append(question_data)
            
        # 計算每科的正確率
        subject_stats = {}
        
        # 先遍歷一次找出這份考卷實際包含的科目
        actual_subjects = set()
        for i, question in enumerate(exam_content['questions']):
            question_number = question['number']
            subject = subjects.get(str(question_number), '未分類')
            actual_subjects.add(subject)
        
        # 只初始化實際出現的科目的統計資料
        for subject in actual_subjects:
            subject_stats[subject] = {
                'total': 0,
                'correct': 0
            }
        
        # 統計每科的總題數和正確題數
        for i, question in enumerate(exam_content['questions']):
            question_number = question['number']
            subject = subjects.get(str(question_number), '未分類')
            subject_stats[subject]['total'] += 1
            
            correct_answer = exam_content['answers'][i] if i < len(exam_content['answers']) else None
            user_answer = None
            possible_keys = [str(i), str(question_number), i, question_number]
            for key in possible_keys:
                if str(key) in user_answers:
                    user_answer = user_answers[str(key)]
                    break
            
            if user_answer == correct_answer:
                subject_stats[subject]['correct'] += 1
        
        # 計算每科的正確率
        subject_percentages = {}
        for subject, stats in subject_stats.items():
            if stats['total'] > 0:
                subject_percentages[subject] = round((stats['correct'] / stats['total']) * 100, 2)
            else:
                subject_percentages[subject] = 0

        # 準備回應數據
        return jsonify({
            'success': True,
            'exam_id': record.exam_id,  # 添加 exam_id
            'questions': questions,
            'subjects': list(actual_subjects),  # 只返回實際出現的科目
            'subject_percentages': subject_percentages  # 只包含實際出現的科目的正確率
        })
        
    except Exception as e:
        print(f"Error in get_exam_detail: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/exam-records/<int:record_id>', methods=['DELETE'])
@handle_errors
def delete_exam_record(record_id):
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'}), 401

    user = User.query.filter_by(username=session['username']).first()
    if not user:
        return jsonify({'success': False, 'error': '用戶不存在'}), 404

    exam_record = ExamRecord.query.get(record_id)
    if not exam_record:
        return jsonify({'success': False, 'error': '考試紀錄不存在'}), 404

    if exam_record.user_id != user.id:
        return jsonify({'success': False, 'error': '無權限刪除此考試紀錄'}), 403

    try:
        # 刪除相關的筆記
        Note.query.filter_by(exam_record_id=record_id).delete()
        # 刪除考試紀錄
        db.session.delete(exam_record)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/questions/<path:filename>')
@handle_errors
def serve_questions(filename):
    """處理靜態文件請求"""
    try:
        # 移除任何可能的路徑操作
        filename = os.path.normpath(filename).replace('\\', '/')
        if filename.startswith('..') or filename.startswith('/'):
            return jsonify({
                'success': False,
                'error': '無效的文件路徑'
            }), 400

        questions_dir = os.path.join(os.path.dirname(__file__), 'questions')
        file_path = os.path.join(questions_dir, filename)

        print(f"Attempting to serve file: {filename}")
        print(f"Full path: {file_path}")

        if not os.path.exists(file_path) or not os.path.isfile(file_path):
            print(f"File not found: {file_path}")
            return jsonify({
                'success': False,
                'error': '文件不存在'
            }), 404

        directory = os.path.dirname(file_path)
        filename = os.path.basename(file_path)

        try:
            return send_from_directory(directory, filename)
        except Exception as e:
            print(f"Error sending file: {str(e)}")
            return jsonify({
                'success': False,
                'error': '無法發送文件'
            }), 500

    except Exception as e:
        print(f"Error serving file {filename}: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'處理文件請求時發生錯誤：{str(e)}'
        }), 500

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))

@app.errorhandler(404)
def not_found_error(error):
    if request.path.startswith('/api/') or request.headers.get('Accept') == 'application/json':
        return jsonify({
            'success': False,
            'error': 'Not found'
        }), 404
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    if request.path.startswith('/api/') or request.headers.get('Accept') == 'application/json':
        return jsonify({
            'success': False,
            'error': 'Internal server error'
        }), 500
    return render_template('500.html'), 500

@app.route('/api/toggle-favorite', methods=['POST'])
@handle_errors
def toggle_favorite():
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'}), 401
        
    data = request.get_json()
    exam_id = data.get('exam_id')
    question_number = data.get('question_number')
    
    if not exam_id or not question_number:
        return jsonify({'success': False, 'error': '缺少必要參數'}), 400
        
    user = User.query.filter_by(username=session['username']).first()
    if not user:
        return jsonify({'success': False, 'error': '用戶不存在'}), 404
        
    existing_favorite = Favorite.query.filter_by(
        user_id=user.id,
        exam_id=exam_id,
        question_number=question_number
    ).first()
    
    if existing_favorite:
        db.session.delete(existing_favorite)
        is_favorite = False
    else:
        new_favorite = Favorite(
            user_id=user.id,
            exam_id=exam_id,
            question_number=question_number
        )
        db.session.add(new_favorite)
        is_favorite = True
        
    db.session.commit()
    return jsonify({
        'success': True,
        'is_favorite': is_favorite
    })

@app.route('/api/get-favorites/<exam_id>', methods=['GET'])
@handle_errors
def get_favorites(exam_id):
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'}), 401
        
    user = User.query.filter_by(username=session['username']).first()
    if not user:
        return jsonify({'success': False, 'error': '用戶不存在'}), 404
        
    from sqlalchemy import cast, Integer
    # 按照題號排序（使用 cast 將字符串轉換為整數）
    favorites = Favorite.query.filter_by(user_id=user.id).order_by(
        cast(Favorite.question_number, Integer).asc()
    ).all()
    
    return jsonify({
        'success': True,
        'favorites': [favorite.question_number for favorite in favorites]
    })

@app.route('/favorites')
def favorites():
    """顯示收藏題目的頁面"""
    if 'username' not in session:
        return redirect(url_for('index'))
    return render_template('favorites.html')

@app.route('/api/get-all-favorites')
@handle_errors
def get_all_favorites():
    """獲取用戶所有收藏題目的 API"""
    if 'username' not in session:
        return jsonify({'success': False, 'error': '請先登入'}), 401
        
    user = User.query.filter_by(username=session['username']).first()
    if not user:
        return jsonify({'success': False, 'error': '用戶不存在'}), 404
        
    favorites = Favorite.query.filter_by(user_id=user.id).order_by(Favorite.created_at.desc()).all()
    
    favorite_questions = []
    for favorite in favorites:
        try:
            # 載入題目內容
            exam_content = load_exam_content(favorite.exam_id)
            if not exam_content or 'questions' not in exam_content or 'answers' not in exam_content:
                continue
                
            # 找到對應的題目
            question = None
            correct_answer = None
            
            for idx, q in enumerate(exam_content['questions']):
                if str(q['number']) == str(favorite.question_number):
                    question = q
                    # 使用相同索引獲取答案
                    correct_answer = exam_content['answers'][idx]
                    break
                    
            if not question or correct_answer is None:
                continue
                
            question_data = {
                'exam_id': favorite.exam_id,
                'question_number': favorite.question_number,
                'content': question['content'],
                'options': question['options'],
                'image': question.get('image'),
                'correct_answer': correct_answer,
                'created_at': favorite.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'exam_title': get_exam_title(favorite.exam_id)
            }
            favorite_questions.append(question_data)
        except Exception as e:
            print(f"Error processing favorite {favorite.exam_id}-{favorite.question_number}: {str(e)}")
            continue
    
    return jsonify({
        'success': True,
        'favorites': favorite_questions
    })

@app.route('/export_exam_word/<int:record_id>')
def export_exam_word(record_id):
    if 'username' not in session:
        return redirect(url_for('login'))

    # 獲取考試記錄
    exam_record = ExamRecord.query.get_or_404(record_id)
    if exam_record.user.username != session['username']:
        abort(403)

    # 獲取考試內容
    exam_content = load_exam_content(exam_record.exam_id)
    if not exam_content:
        abort(404)

    # 獲取筆記
    notes = Note.query.filter_by(exam_record_id=record_id).all()
    print(f"DEBUG: Found {len(notes)} notes for record_id {record_id}")
    
    # 確保 question_number 是整數類型
    notes_dict = {}
    for note in notes:
        print(f"DEBUG: Processing note - question_number: {note.question_number}, content: {note.content}")
        try:
            # 將題號轉換為整數
            q_num = int(note.question_number)
            notes_dict[str(q_num)] = note.content
            print(f"DEBUG: Added note for question {q_num}")
        except (ValueError, TypeError) as e:
            print(f"DEBUG: Error processing note {note.question_number}: {str(e)}")
            continue
    
    print(f"DEBUG: Final notes dictionary: {notes_dict}")
    
    # 創建新的 Word 文件
    doc = Document()
    
    # 設置標題
    title = get_exam_title(exam_record.exam_id)
    heading = doc.add_heading(title, 0)
    heading.alignment = 1  # 置中對齊

    # 添加考試資訊
    doc.add_paragraph(f'考試日期：{exam_record.date.strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'得分：{exam_record.score}')
    doc.add_paragraph(f'作答時間：{exam_record.exam_duration} 秒')
    doc.add_paragraph('') # 空行

    # 遍歷每個題目
    for i, question in enumerate(exam_content['questions']):
        # 題目編號和內容
        q_num = str(question['number'])
        print(f"DEBUG: Processing question {q_num}")
        print(f"DEBUG: Question number type: {type(q_num)}")
        print(f"DEBUG: Notes dict keys: {list(notes_dict.keys())}")
        doc.add_paragraph(f'題目 {q_num}：{question["content"]}', style='List Number')
        
        # 選項
        for j, option in enumerate(question['options']):
            doc.add_paragraph(f'{chr(65+j)}. {option}')
        
        # 答案資訊
        try:
            if exam_record.answers and i < len(exam_record.answers):
                answer_value = exam_record.answers[i]
                if answer_value is not None and isinstance(answer_value, int) and 0 <= answer_value <= 4:
                    user_answer = chr(65 + answer_value)
                else:
                    user_answer = '未作答'
            else:
                user_answer = '未作答'
            
            correct_answer = chr(65 + exam_content['answers'][i]) if i < len(exam_content['answers']) else '未知'
        except (TypeError, IndexError, ValueError):
            user_answer = '未作答'
            correct_answer = '未知'
        
        p = doc.add_paragraph()
        p.add_run('您的答案：').bold = True
        p.add_run(f'{user_answer}    ')
        p.add_run('正確答案：').bold = True
        p.add_run(correct_answer)
        
        # 添加筆記（如果有）
        print(f"DEBUG: Checking notes for question {q_num}")
        print(f"DEBUG: Question number type: {type(q_num)}")
        print(f"DEBUG: Notes dict keys: {notes_dict.keys()}")
        if q_num in notes_dict:
            print(f"DEBUG: Found note for question {q_num}: {notes_dict[q_num]}")
            doc.add_paragraph('') # 空行
            note_para = doc.add_paragraph()
            note_para.style = 'Quote'  # 使用引用樣式
            note_title = note_para.add_run('【筆記】\n')
            note_title.bold = True
            note_title.font.size = Pt(12)
            note_content = note_para.add_run(notes_dict[q_num].strip())
            note_content.font.size = Pt(11)
        else:
            print(f"DEBUG: No note found for question {q_num}")
        
        # 添加分隔線
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('') # 空行

    # 將文件保存到記憶體中
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return send_file(
        docx_file,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{title}.docx'
    )

@app.route('/admin/exam_info/<exam_id>')
def admin_exam_info(exam_id):
    if not session.get('is_admin'):
        abort(403)
    
    exam_content = load_exam_content(exam_id)
    if not exam_content:
        abort(404)
    
    # 獲取所有用戶對這份考卷的筆記
    notes = Note.query.join(ExamRecord).filter(
        ExamRecord.exam_id == exam_id
    ).all()
    
    notes_by_question = {}
    for note in notes:
        if note.question_number not in notes_by_question:
            notes_by_question[note.question_number] = []
        notes_by_question[note.question_number].append({
            'username': note.user.username,
            'content': note.content,
            'created_at': note.created_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'exam_content': exam_content,
        'notes': notes_by_question
    })

@app.route('/admin/exam_list')
def admin_exam_list():
    if not session.get('is_admin'):
        abort(403)
    
    exam_years = ['110', '111', '112', '113']
    exam_list = []
    for year in exam_years:
        exam_list.extend([
            {'id': f'{year}_1', 'name': f'{year}年度前80題'},
            {'id': f'{year}_2', 'name': f'{year}年度後80題'}
        ])
    
    return jsonify(exam_list)

@app.route('/api/exam_notes/<exam_id>/<int:question_number>')
def get_question_notes(exam_id, question_number):
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    try:
        user = User.query.filter_by(username=session['username']).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404

        # 找到該用戶對應考試的所有記錄
        exam_records = ExamRecord.query.filter_by(
            user_id=user.id,
            exam_id=exam_id
        ).all()

        # 收集所有筆記
        notes = []
        for record in exam_records:
            note = Note.query.filter_by(
                exam_record_id=record.id,
                question_number=question_number
            ).first()
            if note:
                notes.append(note.to_dict())

        return jsonify({
            'success': True,
            'notes': notes
        })

    except Exception as e:
        print(f"Error getting notes: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
